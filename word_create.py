from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# إنشاء ملف Word جديد
doc = Document()

# بيانات العينة (مثال Aspartame)
sample_info = {
    "Sample Name": "Aspartame",
    "Sample Type": "Unknown",
    "Injection #": "1",
    "Injection Volume": "20 µL",
    "Run Time": "15.0 Minutes",
    "Date Acquired": "06-Jul-2023 09:41:46 AM",
    "Date Processed": "06-Jul-2023 11:37:50 AM",
    "Acquired By": "System",
    "Sample Set Name": "Asp_Set1",
    "Acq. Method Set": "Asp_M",
    "Processing Method": "Aspartame",
    "Channel Name": "2487Channel 1",
    "Proc. Chnl. Descr.": "2487Channel 1",
}

# بيانات Peaks (مثال)
peak_table = [
    ["Aspartame", "9.50", "1000234", "99.93", "297438", "7572.38"],
]

# USP Tailing (مثال)
usp_tailing = [["1", "1.057"]]

# --- حقنات 1 إلى 6 ---
for inj in range(1, 7):
    doc.add_heading(f"Injection {inj}", level=1)
    
    # Sample Information
    table = doc.add_table(rows=0, cols=2)
    for k, v in sample_info.items():
        row = table.add_row().cells
        row[0].text = k
        row[1].text = str(inj) if k == "Injection #" else v
    doc.add_paragraph("")

    # Placeholder للكروماتوجرام
    p = doc.add_paragraph("[ Chromatogram Image Placeholder ]")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")
    
    # Peaks Table
    doc.add_heading("Peaks", level=2)
    peak_t = doc.add_table(rows=1, cols=6)
    hdr = peak_t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text, hdr[5].text = \
        ["Peak Name", "RT", "Area", "% Area", "Height", "USP Plate Count"]
    for row in peak_table:
        r = peak_t.add_row().cells
        for i, val in enumerate(row):
            r[i].text = val
    
    # USP Tailing
    doc.add_heading("USP Tailing", level=2)
    tail_t = doc.add_table(rows=1, cols=2)
    tail_t.rows[0].cells[0].text, tail_t.rows[0].cells[1].text = ["Peak", "USP Tailing"]
    for row in usp_tailing:
        r = tail_t.add_row().cells
        r[0].text, r[1].text = row
    
    doc.add_page_break()

# --- System Precision Overlay ---
doc.add_heading("System Precision Overlay", level=1)
doc.add_paragraph("[ Overlay Chromatogram Placeholder ]").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_page_break()

# --- Comprehensive Analysis ---
doc.add_heading("Comprehensive Analysis", level=1)
analysis_t = doc.add_table(rows=1, cols=4)
hdr = analysis_t.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = ["Parameter", "Result", "Specification", "Status"]
rows = [
    ["Chromatographic Purity", "99.93%", "≥ 98.0%", "PASS"],
    ["5-Benzyl-3,6-dioxo-2-piperazineacetic Acid", "0.07%", "≤ 1.0%", "PASS"],
    ["System Suitability", "All Parameters", "USP-NF 2024", "PASS"],
]
for row in rows:
    r = analysis_t.add_row().cells
    for i, val in enumerate(row):
        r[i].text = val

doc.add_paragraph("[ Bar Chart Placeholder ]").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_page_break()

# --- Conclusion & Approval ---
doc.add_heading("Conclusion & Approval", level=1)
doc.add_paragraph("Batch Approved for Release\nAll quality parameters meet USP-NF 2024 requirements.")
doc.add_paragraph("")
doc.add_paragraph("QC Analyst: Dr. Naema Lofty\n(Signature)\nDigitally Signed & Verified")
doc.add_paragraph("QC Manager: Dr. Montaser Elzayaat\n(Signature)\nDigitally Signed & Verified")

# حفظ الملف
doc.save("HPLC_Report_Aspartame.docx")
print("Report saved as HPLC_Report_Aspartame.docx")
import os
print("Saved to:", os.path.abspath("HPLC_Report_Aspartame.docx"))
